"""
report.py — DOCX-отчёт с правильными названиями.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import database as db
from config import DATA_DIR


def generate_report(meeting_id):
    mt = db.get_meeting(meeting_id)
    if not mt:
        return None

    label = db.type_label(mt["type"])
    member_map = {m["pin"]: m["name"] for m in db.get_members()}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Заголовок
    t = doc.add_heading(f"РЕЗУЛЬТАТЫ ТАЙНОГО ГОЛОСОВАНИЯ", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{label} №{mt['protocol_number']}").bold = True

    doc.add_paragraph("")

    # Инфо
    info = doc.add_paragraph()
    info.add_run("Тип: ").bold = True
    info.add_run(f"{label}\n")
    info.add_run("Дата: ").bold = True
    info.add_run(f"{mt['date']}\n")
    info.add_run("Протокол №: ").bold = True
    info.add_run(f"{mt['protocol_number']}\n")
    info.add_run("Присутствовало: ").bold = True
    info.add_run(f"{len(mt['attendees'])} чел.")

    # Присутствующие
    doc.add_heading("Присутствующие:", level=2)
    names = sorted(member_map.get(p, p) for p in mt["attendees"])
    for i, name in enumerate(names, 1):
        doc.add_paragraph(f"{i}. {name}", style="List Number")

    doc.add_paragraph("")

    # Вопросы
    closed_q = [q for q in mt["questions"] if q["status"] == "closed"]
    if closed_q:
        doc.add_heading("Результаты голосования:", level=2)

        for i, q in enumerate(closed_q, 1):
            v = q["votes"]
            total = v["for"] + v["against"] + v["abstain"]

            p = doc.add_paragraph()
            p.add_run(f"Вопрос {i}: ").bold = True
            p.add_run(q["text"])

            # Таблица
            table = doc.add_table(rows=2, cols=4)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(["", "За", "Против", "Воздержался"]):
                c = table.rows[0].cells[j]
                c.text = h
                for par in c.paragraphs:
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in par.runs:
                        r.bold = True

            row = table.rows[1]
            row.cells[0].text = "Голосов"
            row.cells[1].text = str(v["for"])
            row.cells[2].text = str(v["against"])
            row.cells[3].text = str(v["abstain"])
            for j in range(4):
                for par in row.cells[j].paragraphs:
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

            rp = doc.add_paragraph()
            rp.add_run(f"Всего голосов: {total} из {len(mt['attendees'])}").italic = True
            rp.add_run("\n")

            if v["for"] > v["against"]:
                run = rp.add_run("РЕШЕНИЕ: ПРИНЯТО")
                run.bold = True
                run.font.color.rgb = RGBColor(0x1B, 0x7A, 0x2B)
            elif v["for"] < v["against"]:
                run = rp.add_run("РЕШЕНИЕ: НЕ ПРИНЯТО")
                run.bold = True
                run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
            else:
                run = rp.add_run("РЕШЕНИЕ: ГОЛОСА РАВНЫ")
                run.bold = True

            doc.add_paragraph("")

    # Подписи
    doc.add_paragraph("")
    s1 = doc.add_paragraph()
    if mt["type"] == "seminar":
        s1.add_run("Руководитель семинара: ").bold = True
    else:
        s1.add_run("Председатель учёного совета: ").bold = True
    s1.add_run("_" * 30)

    s2 = doc.add_paragraph()
    s2.add_run("Секретарь: ").bold = True
    s2.add_run("_" * 30)

    # Сохранить
    os.makedirs(DATA_DIR, exist_ok=True)
    safe = mt["date"].replace(".", "-").replace("/", "-")
    prefix = "seminar" if mt["type"] == "seminar" else "council"
    fn = f"{prefix}_{mt['protocol_number']}_{safe}.docx"
    fp = os.path.join(DATA_DIR, fn)
    doc.save(fp)
    return fp
